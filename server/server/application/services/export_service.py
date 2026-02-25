"""
Export service for document conversion.

Article III (TDD): Service will be tested via pytest
Article V (Documentation): All methods documented
"""

from typing import Protocol
from io import BytesIO


class ExportFormat(Protocol):
    """Protocol for export format handlers."""

    def export(self, content: str, title: str = "Document") -> bytes:
        """Export content to specific format.

        Args:
            content: Markdown content to export
            title: Document title

        Returns:
            Exported content as bytes
        """
        ...


class PDFExportService:
    """Service for exporting documents to PDF format.

    Uses reportlab (pure Python) for PDF generation.
    Converts markdown to styled PDF with proper formatting.
    """

    def __init__(self) -> None:
        """Initialize PDF export service."""
        pass

    def export_to_pdf(self, markdown_content: str, title: str = "Document") -> bytes:
        """Export markdown content to PDF.

        Args:
            markdown_content: Markdown text to convert
            title: Document title for header

        Returns:
            PDF bytes

        Raises:
            ImportError: If required libraries not installed
            Exception: If conversion fails
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_LEFT
            import markdown
            from html.parser import HTMLParser
        except ImportError as e:
            raise ImportError(
                "PDF export requires 'reportlab' and 'markdown'. "
                "Install with: poetry add reportlab markdown"
            ) from e

        # Create PDF document
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm,
        )

        # Setup styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
        ))
        styles.add(ParagraphStyle(
            name='CustomHeading2',
            parent=styles['Heading2'],
            fontSize=18,
            spaceBefore=20,
            spaceAfter=10,
        ))
        styles.add(ParagraphStyle(
            name='CustomHeading3',
            parent=styles['Heading3'],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=8,
        ))
        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=10,
        ))

        # Build content
        story = []
        story.append(Paragraph(title, styles['CustomTitle']))
        story.append(Spacer(1, 12))

        # Convert markdown to HTML
        html_content = markdown.markdown(
            markdown_content,
            extensions=["extra", "toc"]
        )

        # Parse HTML and convert to reportlab elements
        class HTMLToReportlab(HTMLParser):
            def __init__(self):
                super().__init__()
                self.current_text = []
                self.in_paragraph = False
                self.in_heading = False
                self.heading_level = 0

            def handle_starttag(self, tag, attrs):
                if tag == 'p':
                    self.in_paragraph = True
                    self.current_text = []
                elif tag in ['h1', 'h2', 'h3']:
                    self.in_heading = True
                    self.heading_level = int(tag[1])
                    self.current_text = []
                elif tag == 'strong':
                    self.current_text.append('<b>')
                elif tag == 'em':
                    self.current_text.append('<i>')
                elif tag == 'code':
                    self.current_text.append('<font face="Courier">')

            def handle_endtag(self, tag):
                if tag == 'p':
                    self.in_paragraph = False
                    text = ''.join(self.current_text)
                    if text.strip():
                        story.append(Paragraph(text, styles['CustomBody']))
                    self.current_text = []
                elif tag in ['h1', 'h2', 'h3']:
                    self.in_heading = False
                    text = ''.join(self.current_text)
                    if text.strip():
                        if self.heading_level == 1:
                            story.append(Paragraph(text, styles['CustomHeading2']))
                        elif self.heading_level == 2:
                            story.append(Paragraph(text, styles['CustomHeading2']))
                        else:
                            story.append(Paragraph(text, styles['CustomHeading3']))
                    self.current_text = []
                elif tag == 'strong':
                    self.current_text.append('</b>')
                elif tag == 'em':
                    self.current_text.append('</i>')
                elif tag == 'code':
                    self.current_text.append('</font>')

            def handle_data(self, data):
                if self.in_paragraph or self.in_heading:
                    self.current_text.append(data)

        parser = HTMLToReportlab()
        parser.feed(html_content)

        # Build PDF
        doc.build(story)

        # Get bytes
        buffer.seek(0)
        return buffer.read()


class DOCXExportService:
    """Service for exporting documents to DOCX format.

    Uses python-docx library for Word document generation.
    """

    def export_to_docx(self, markdown_content: str, title: str = "Document") -> bytes:
        """Export markdown content to DOCX.

        Args:
            markdown_content: Markdown text to convert
            title: Document title

        Returns:
            DOCX bytes

        Raises:
            ImportError: If python-docx not installed
        """
        try:
            from docx import Document
            import markdown
            from html.parser import HTMLParser
        except ImportError as e:
            raise ImportError(
                "DOCX export requires 'python-docx' and 'markdown'. "
                "Install with: poetry add python-docx markdown"
            ) from e

        # Convert markdown to HTML first
        html_content = markdown.markdown(
            markdown_content,
            extensions=["extra"]
        )

        # Create DOCX document
        doc = Document()
        doc.add_heading(title, 0)

        # Simple HTML to DOCX conversion
        class HTMLToDocx(HTMLParser):
            def __init__(self, document):
                super().__init__()
                self.doc = document
                self.current_paragraph = None

            def handle_starttag(self, tag, attrs):
                if tag in ['h1', 'h2', 'h3']:
                    level = int(tag[1])
                    self.current_paragraph = self.doc.add_heading('', level=level)
                elif tag == 'p':
                    self.current_paragraph = self.doc.add_paragraph()
                elif tag == 'code':
                    pass  # Will be handled in handle_data

            def handle_data(self, data):
                if self.current_paragraph is not None:
                    self.current_paragraph.add_run(data)

        parser = HTMLToDocx(doc)
        parser.feed(html_content)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.read()


# Convenience function for simple usage
def export_to_pdf(content: str, title: str = "Document") -> bytes:
    """Convenience function for PDF export.

    Args:
        content: Markdown content
        title: Document title

    Returns:
        PDF bytes
    """
    service = PDFExportService()
    return service.export_to_pdf(content, title)
